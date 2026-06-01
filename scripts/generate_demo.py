#!/usr/bin/env python3
"""
ThinkVault Demo 数据生成器

生成 5 个不同类型的示例文档到 demo_docs/ 目录，用于演示 ThinkVault 的
多格式解析能力。

生成文件：
  demo_docs/
  ├── 01_技术架构说明.txt        — 技术文档 (TXT)
  ├── 02_ThinkVault项目说明.md   — 项目说明 (Markdown)
  ├── 03_周例会会议纪要.txt      — 会议纪要 (TXT)
  ├── 04_产品规格说明书.docx     — 产品规格 (DOCX)
  └── 05_销售数据统计.xlsx       — 数据表格 (XLSX)

用法：
  python scripts/generate_demo.py [--output demo_docs]

导入 ThinkVault：
  启动服务后，在 Web UI 上传 demo_docs/ 目录下的文件，
  或通过 API:
    curl -F "file=@demo_docs/01_技术架构说明.txt" http://localhost:8000/api/documents/upload
"""

import argparse
import sys
from pathlib import Path

# 确保项目根目录在 sys.path
_SCRIPT_DIR = Path(__file__).resolve().parent
_PROJECT_ROOT = _SCRIPT_DIR.parent
sys.path.insert(0, str(_PROJECT_ROOT))

OUTPUT_DIR = _PROJECT_ROOT / "demo_docs"


def generate_txt_tech_arch() -> Path:
    """生成技术架构说明文档 (TXT)"""
    content = """\
ThinkVault 技术架构说明
========================

1. 系统概述
-----------
ThinkVault 是一款运行在本地的 RAG（检索增强生成）知识库系统。
它允许用户将 PDF、DOCX、TXT、PPTX、XLSX 等文档导入知识库，
并通过自然语言对话进行检索和问答。

2. 核心架构
-----------
系统采用模块化设计，主要包含以下组件：

2.1 文档解析层 (Parser)
  - 支持格式：PDF / DOCX / PPTX / XLSX / TXT / Markdown
  - 底层依赖：PyMuPDF (PDF)、python-docx (DOCX)、python-pptx (PPTX)、openpyxl (XLSX)
  - 统一输出：ParsedDocument 数据结构，包含段落、表格、原始文本

2.2 文本分块层 (Chunker)
  - 算法：固定窗口 + 重叠分块
  - 默认参数：chunk_size=512, chunk_overlap=64
  - 策略：优先在段落边界切分，保证语义完整性

2.3 向量化层 (Embedder)
  - 模型：BAAI/bge-small-zh-v1.5 (384 维)
  - 框架：sentence-transformers
  - 支持 GPU 加速（CUDA / MPS）

2.4 检索引擎 (Retriever)
  - BM25 关键字检索 + 向量语义检索
  - Cross-encoder 重排序（ms-marco-MiniLM-L-6-v2）
  - 智能意图判断（关键词 + 语义相似度）

2.5 LLM 推理层 (ThinkVaultLLM)
  - 框架：httpx + OpenAI 兼容 API（对接 Ollama）
  - 格式：通过 Ollama 拉取模型（如 ollama pull llama3.2:3b）
  - 推荐模型：llama3.2:3b

2.6 API 服务层 (FastAPI)
  - 对话接口：/api/chat (POST), /api/chat/stream (SSE)
  - 文档管理：/api/documents (CRUD)
  - 模型管理：/api/model (状态/加载/卸载)
  - 会话持久化：SQLite 存储

3. 数据流
---------
  用户上传文档 -> Parser 解析 -> Chunker 分块 -> Embedder 向量化 -> VectorStore 存储
  用户提问 -> Retriever 检索 -> LLM 推理 -> SSE 流式返回答案

4. 部署方式
-----------
  - 本地 Python 环境：python -m thinkvault
  - Docker Compose：docker-compose up -d
  - 环境变量配置：THINKVAULT_API_TOKEN, THINKVAULT_CORS_ORIGINS 等

5. 性能指标
-----------
  - 文档解析速度：~2.5 MB/s (PDF)
  - 嵌入速度：~35 文本块/秒 (CPU)，~280 文本块/秒 (GPU)
  - 检索延迟：<200ms (10,000 文档以内)
  - 推理速度：~15 tokens/秒 (Llama-3.2-3B, CPU)
"""
    path = OUTPUT_DIR / "01_技术架构说明.txt"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")
    print(f"  [OK] {path.name}")
    return path


def generate_md_project_intro() -> Path:
    """生成项目说明文档 (Markdown)"""
    content = """\
# ThinkVault 项目说明

## 项目背景

在日常工作中，我们积累了大量的技术文档、会议纪要、产品规格书和数据分析报告。
传统的关键词搜索无法理解语义，而商业知识库产品往往需要将数据上传至云端，
存在隐私泄露风险。

**ThinkVault** 应运而生 — 一款完全本地运行的智能知识库系统。

## 核心功能

### 1. 多格式文档支持

| 格式 | 扩展名 | 解析引擎 |
|------|--------|----------|
| PDF 文档 | .pdf | PyMuPDF |
| Word 文档 | .docx | python-docx |
| PowerPoint 演示 | .pptx | python-pptx |
| Excel 表格 | .xlsx / .xlsm | openpyxl |
| 纯文本 | .txt | 内置 |
| Markdown | .md | 内置 |

### 2. 智能检索

- **混合检索**：BM25 关键字匹配 + 向量语义检索，取长补短
- **重排序**：Cross-encoder 对候选结果精细打分
- **意图识别**：自动判断用户输入是否需要检索知识库

### 3. 对话管理

- 多会话支持（创建、重命名、删除）
- 消息持久化（SQLite），重启不丢失
- SSE 流式输出，逐 Token 返回

### 4. 安全与隐私

- 全本地运行，数据不出机器
- API Token 认证（可选）
- 基于 IP 的速率限制
- CORS 白名单控制

## 快速开始

```bash
# 安装依赖
pip install -r requirements.txt

# 启动服务
python -m thinkvault

# 浏览器访问
open http://localhost:8000
```

## 技术栈

- **后端**：Python 3.10+, FastAPI, Uvicorn
- **推理后端**：Ollama (OpenAI 兼容), BGE (sentence-transformers)
- **存储**：ChromaDB (向量), SQLite (元数据)
- **前端**：原生 HTML/CSS/JS (无框架依赖)

## 许可证

MIT License - 自由使用、修改和分发。
"""
    path = OUTPUT_DIR / "02_ThinkVault项目说明.md"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")
    print(f"  [OK] {path.name}")
    return path


def generate_txt_meeting_minutes() -> Path:
    """生成会议纪要 (TXT)"""
    content = """\
ThinkVault 项目周例会会议纪要
================================

会议时间：2026年5月28日 14:00-15:30
会议地点：线上腾讯会议
参会人员：张三（项目经理）、李四（后端开发）、王五（前端开发）、赵六（测试）

一、上周工作回顾
-----------------

1. 后端开发（李四）
   - 完成 BM25 + Cross-encoder 混合检索功能开发
   - 修复 PDF 多页表格提取 Bug
   - 新增文档上传 API 的格式校验

2. 前端开发（王五）
   - 优化 SSE 流式消息展示，支持 Markdown 渲染
   - 添加文档上传进度条
   - 修复移动端响应式布局问题

3. 测试（赵六）
   - 编写 V2.0 集成测试 19 个用例
   - 完成 PPT/Excel 解析的 Mock 测试
   - 发现 DOCX 空文档处理缺失问题（已记录）

二、本周计划
--------------

1. 张三：协调 Docker 镜像发布流程，编写 Release Notes
2. 李四：
   - 完成 Dockerfile 和 docker-compose.yml
   - 修复 DOCX 空文档 parse_error 设置
   - 添加 Cross-encoder 首次下载的友好提示

3. 王五：
   - 实现知识库列表分组视图
   - 添加对话搜索功能

4. 赵六：
   - 编写集成测试自动化脚本（启动服务 → 运行测试 → 关闭 → 报告）
   - 全量回归测试

三、风险与阻塞项
-----------------

1. Ollama 默认监听端口 11434，需确保防火墙允许
   -> 已记录在 README 中，暂不影响 V2.0 发布

2. ChromaDB 在 Windows 上的权限问题
   -> 通过设置 THINKVAULT_DATA_DIR 环境变量解决

四、里程碑
-----------

- V1.0 内部测试版：2026-04-15（已完成）
- V2.0 发布就绪：2026-06-01（目标）
- V2.1 功能增强：2026-06-15（计划中）

五、决议事项
---------------

1. 确定 V2.0 采用 MIT 开源协议
2. Docker 镜像基础选型：python:3.10-slim（CPU）/ nvidia/cuda:12.1-runtime（GPU）
3. 首次发布暂不包含 OCR 功能，列入 V2.1 计划
"""
    path = OUTPUT_DIR / "03_周例会会议纪要.txt"
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(content, encoding="utf-8")
    print(f"  [OK] {path.name}")
    return path


def generate_docx_product_spec() -> Path:
    """生成产品规格说明书 (DOCX)"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        print("  [SKIP] python-docx 未安装，跳过 DOCX 生成")
        return None

    doc = Document()

    # 标题
    title = doc.add_heading("ThinkVault 产品规格说明书", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("版本：V2.0 | 日期：2026-06-01 | 状态：发布就绪")

    # 1. 产品概述
    doc.add_heading("1. 产品概述", level=1)
    doc.add_paragraph(
        "ThinkVault 是一款运行在本地的个人 AI 知识库系统。"
        "它能够将用户的各种文档（PDF、DOCX、PPTX、XLSX、TXT 等）"
        "转化为可检索的知识库，并通过自然语言对话进行问答。"
    )

    # 2. 功能规格
    doc.add_heading("2. 功能规格", level=1)

    doc.add_heading("2.1 文档解析", level=2)

    # 添加表格
    table = doc.add_table(rows=7, cols=4, style="Light Grid Accent 1")
    headers = ["格式", "扩展名", "解析引擎", "说明"]
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h

    data = [
        ["PDF", ".pdf", "PyMuPDF", "支持文本提取、表格识别"],
        ["Word", ".docx", "python-docx", "支持段落和表格提取"],
        ["PPT", ".pptx", "python-pptx", "支持文本和表格提取"],
        ["Excel", ".xlsx/.xlsm", "openpyxl", "多 Sheet、空行过滤"],
        ["文本", ".txt", "内置", "UTF-8/GBK 自动检测"],
        ["Markdown", ".md", "内置", "同 TXT 处理"],
    ]
    for r, row_data in enumerate(data, start=1):
        for c, val in enumerate(row_data):
            table.rows[r].cells[c].text = val

    doc.add_heading("2.2 智能检索", level=2)
    doc.add_paragraph("混合检索策略：BM25 关键字匹配 + 向量语义检索", style="List Bullet")
    doc.add_paragraph("Cross-encoder 重排序：ms-marco-MiniLM-L-6-v2", style="List Bullet")
    doc.add_paragraph("检索意图智能判断：关键词 + 语义相似度双级判断", style="List Bullet")

    doc.add_heading("2.3 对话管理", level=2)
    doc.add_paragraph("多会话创建、重命名、删除", style="List Bullet")
    doc.add_paragraph("消息持久化（SQLite），重启不丢失", style="List Bullet")
    doc.add_paragraph("SSE 流式输出，逐 Token 返回", style="List Bullet")

    # 3. 技术规格
    doc.add_heading("3. 技术规格", level=1)

    tech_table = doc.add_table(rows=9, cols=2, style="Light Grid Accent 1")
    tech_data = [
        ("运行环境", "Python 3.10+"),
        ("Web 框架", "FastAPI + Uvicorn"),
        ("嵌入模型", "BAAI/bge-small-zh-v1.5 (384维)"),
        ("推理后端", "Ollama (OpenAI 兼容)"),
        ("向量数据库", "ChromaDB"),
        ("元数据存储", "SQLite"),
        ("前端技术", "原生 HTML/CSS/JS"),
        ("容器化", "Docker + Docker Compose"),
    ]
    for r, (k, v) in enumerate(tech_data):
        tech_table.rows[r].cells[0].text = k
        tech_table.rows[r + 1 if r == 0 else r].cells[1].text = v

    # 修正表格行
    tech_table = doc.add_table(rows=len(tech_data), cols=2, style="Light Grid Accent 1")
    for r, (k, v) in enumerate(tech_data):
        tech_table.rows[r].cells[0].text = k
        tech_table.rows[r].cells[1].text = v

    # 4. 部署规格
    doc.add_heading("4. 部署规格", level=1)
    doc.add_paragraph("最低配置：4 核 CPU / 8GB RAM / 10GB 磁盘", style="List Bullet")
    doc.add_paragraph("推荐配置：8 核 CPU / 16GB RAM / 50GB SSD", style="List Bullet")
    doc.add_paragraph("支持 Docker Compose 一键部署", style="List Bullet")
    doc.add_paragraph("GPU 加速：CUDA 12.1+，通过 nvidia-container-toolkit", style="List Bullet")

    path = OUTPUT_DIR / "04_产品规格说明书.docx"
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    print(f"  [OK] {path.name}")
    return path


def generate_xlsx_sales_data() -> Path:
    """生成销售数据统计表格 (XLSX)"""
    try:
        import openpyxl
        from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
        from openpyxl.utils import get_column_letter
    except ImportError:
        print("  [SKIP] openpyxl 未安装，跳过 XLSX 生成")
        return None

    wb = openpyxl.Workbook()

    # ---- Sheet 1: 月度销售汇总 ----
    ws1 = wb.active
    ws1.title = "月度销售汇总"

    headers = ["月份", "产品名称", "销售数量", "单价(元)", "销售额(元)", "区域", "销售人员"]
    header_font = Font(bold=True, color="FFFFFF", size=11)
    header_fill = PatternFill(start_color="4472C4", end_color="4472C4", fill_type="solid")
    header_align = Alignment(horizontal="center", vertical="center")

    for col, h in enumerate(headers, start=1):
        cell = ws1.cell(row=1, column=col, value=h)
        cell.font = header_font
        cell.fill = header_fill
        cell.alignment = header_align

    sales_data = [
        ["2026-01", "ThinkVault Pro", 45, 299, 13455, "华东", "张三"],
        ["2026-01", "ThinkVault Lite", 120, 99, 11880, "华东", "李四"],
        ["2026-02", "ThinkVault Pro", 52, 299, 15548, "华南", "王五"],
        ["2026-02", "ThinkVault Lite", 135, 99, 13365, "华南", "赵六"],
        ["2026-03", "ThinkVault Pro", 68, 299, 20332, "华北", "张三"],
        ["2026-03", "ThinkVault Lite", 160, 99, 15840, "华北", "钱七"],
        ["2026-04", "ThinkVault Pro", 73, 299, 21827, "华东", "李四"],
        ["2026-04", "ThinkVault Lite", 185, 99, 18315, "华东", "孙八"],
        ["2026-05", "ThinkVault Pro", 90, 299, 26910, "华南", "王五"],
        ["2026-05", "ThinkVault Lite", 210, 99, 20790, "华南", "周九"],
    ]

    for r, row_data in enumerate(sales_data, start=2):
        for c, val in enumerate(row_data, start=1):
            ws1.cell(row=r, column=c, value=val)

    # 自动调整列宽
    for col in range(1, len(headers) + 1):
        ws1.column_dimensions[get_column_letter(col)].width = 15

    # ---- Sheet 2: 产品销售排行 ----
    ws2 = wb.create_sheet("产品销售排行")

    rank_headers = ["排名", "产品名称", "累计销量", "累计销售额(元)", "平均单价(元)"]
    for col, h in enumerate(rank_headers, start=1):
        cell = ws2.cell(row=1, column=col, value=h)
        cell.font = header_font
        cell.fill = PatternFill(start_color="548235", end_color="548235", fill_type="solid")
        cell.alignment = header_align

    rank_data = [
        [1, "ThinkVault Lite", 810, 80190, 99],
        [2, "ThinkVault Pro", 328, 98072, 299],
        [3, "ThinkVault Enterprise", 45, 44955, 999],
    ]

    for r, row_data in enumerate(rank_data, start=2):
        for c, val in enumerate(row_data, start=1):
            ws2.cell(row=r, column=c, value=val)

    for col in range(1, len(rank_headers) + 1):
        ws2.column_dimensions[get_column_letter(col)].width = 18

    path = OUTPUT_DIR / "05_销售数据统计.xlsx"
    path.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(path))
    print(f"  [OK] {path.name}")
    return path


def main():
    global OUTPUT_DIR
    parser = argparse.ArgumentParser(description="ThinkVault Demo 数据生成器")
    parser.add_argument("--output", default=str(OUTPUT_DIR), help=f"输出目录 (默认: {OUTPUT_DIR})")
    args = parser.parse_args()

    OUTPUT_DIR = Path(args.output)
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    print(f"ThinkVault Demo 数据生成器")
    print(f"输出目录: {OUTPUT_DIR}\n")

    results = []
    results.append(generate_txt_tech_arch())
    results.append(generate_md_project_intro())
    results.append(generate_txt_meeting_minutes())
    results.append(generate_docx_product_spec())
    results.append(generate_xlsx_sales_data())

    generated = [r for r in results if r is not None]
    print(f"\n生成完成：共 {len(generated)} 个文件")
    print(f"\n导入方式：")
    print(f"  1. Web UI: 访问 http://localhost:8000 → 上传文档")
    print(f"  2. API:")
    for f in generated:
        print(f"     curl -F \"file=@{f}\" http://localhost:8000/api/documents/upload")
    print(f"  3. 拖拽导入: 将文件直接拖入 Web UI 上传区")


if __name__ == "__main__":
    main()